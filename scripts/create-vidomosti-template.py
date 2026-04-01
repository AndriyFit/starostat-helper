"""Generate .docx template: Відомості про зареєстрованих у житловому приміщенні осіб"""
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

def set_cell_border(cell, **kwargs):
    """Set cell border — None removes borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tcBorders>')
    tcPr.append(tcBorders)

def remove_table_borders(table):
    """Remove all borders from a table."""
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)

def add_run(paragraph, text, bold=False, size=None, font_name='Times New Roman'):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return run

# === CORNER STAMP (2-column table, no borders) ===
header_table = doc.add_table(rows=1, cols=2)
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
for cell in header_table.columns[0].cells:
    cell.width = Cm(10)
for cell in header_table.columns[1].cells:
    cell.width = Cm(7)

# Left cell — stamp
left_cell = header_table.cell(0, 0)
left_cell.text = ''
p = left_cell.paragraphs[0]
add_run(p, 'Жовтанецька сільська рада', bold=True, size=12)

p = left_cell.add_paragraph()
add_run(p, '80431, Львівська область, Львівський район,', size=9)
p = left_cell.add_paragraph()
add_run(p, 'село Жовтанці, вулиця Львівська, 2', size=9)

p = left_cell.add_paragraph()
p.space_before = Pt(6)
add_run(p, '№ _____ від ', size=11)
add_run(p, '{{date}}', bold=True, size=11)

p = left_cell.add_paragraph()
add_run(p, 'на', size=11)
p = left_cell.add_paragraph()
add_run(p, '№ _____ від _____', size=11)

# Right cell — "За запитом"
right_cell = header_table.cell(0, 1)
right_cell.text = ''
p = right_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p, 'За запитом:', size=11)

p = right_cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p, '{{PIB}}', bold=True, size=12)

remove_table_borders(header_table)

# === SPACER ===
doc.add_paragraph()
doc.add_paragraph()

# === TITLE ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Відомості', bold=True, size=14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'про зареєстрованих у житловому приміщенні осіб', bold=True, size=14)

doc.add_paragraph()

# === BODY ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_run(p, 'Видана      ', size=12)
add_run(p, 'власнику', size=12)
add_run(p, '      житлового приміщення', size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '{{PIB}}', bold=True, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '(прізвище, ім\'я, по батькові)', size=9)

doc.add_paragraph()

p = doc.add_paragraph()
add_run(p, 'про те, що за адресою:', size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '{{address}}', bold=True, size=12)

doc.add_paragraph()
p = doc.add_paragraph()
add_run(p, 'зареєстровані:', size=12)

doc.add_paragraph()

# === MEMBERS TABLE ===
members_table = doc.add_table(rows=2, cols=4)
members_table.alignment = WD_TABLE_ALIGNMENT.CENTER
members_table.style = 'Table Grid'

# Header row
headers = ['№', 'Прізвище, ім\'я, по батькові', 'Дата\nнародження', 'Період реєстрації']
for i, h in enumerate(headers):
    cell = members_table.cell(0, i)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, bold=True, size=11)

# Data row with placeholder
members_table.cell(1, 0).text = ''
p = members_table.cell(1, 0).paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Use multiline text placeholder for members
members_table.cell(1, 1).text = ''
members_table.cell(1, 2).text = ''
members_table.cell(1, 3).text = ''

# Remove data row — we'll use text-based members instead
# (docxtemplater can't easily loop table rows without special syntax)
# Remove the data row
members_table._tbl.remove(members_table.rows[1]._tr)

doc.add_paragraph()

# Members as text (simpler approach compatible with docxtemplater)
p = doc.add_paragraph()
add_run(p, '{{members}}', size=12)

doc.add_paragraph()

p = doc.add_paragraph()
add_run(p, 'Усього ', size=12)
add_run(p, '{{total}}', bold=True, size=12)
add_run(p, ' осіб', size=12)

# === SPACER ===
doc.add_paragraph()
doc.add_paragraph()

# === SIGNATURE ===
sig_table = doc.add_table(rows=1, cols=3)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

sig_table.cell(0, 0).text = ''
p = sig_table.cell(0, 0).paragraphs[0]
add_run(p, 'Староста Ременівського\nстаростинського округу', bold=True, size=11)

sig_table.cell(0, 1).text = ''
p = sig_table.cell(0, 1).paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '(підпис)', size=9)

sig_table.cell(0, 2).text = ''
p = sig_table.cell(0, 2).paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p, '{{starosta}}', bold=True, size=11)

# Add hint under signature
sig_table2 = doc.add_table(rows=1, cols=3)
sig_table2.alignment = WD_TABLE_ALIGNMENT.CENTER

sig_table2.cell(0, 0).text = ''
p = sig_table2.cell(0, 0).paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '(посада)', size=9)

sig_table2.cell(0, 1).text = ''
p = sig_table2.cell(0, 1).paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '', size=9)

sig_table2.cell(0, 2).text = ''
p = sig_table2.cell(0, 2).paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p, '(прізвище, ініціали)', size=9)

remove_table_borders(sig_table)
remove_table_borders(sig_table2)

# Save
out = os.path.join(os.path.dirname(__file__), '..', 'kb', 'templates', 'vidomosti-zareyestrovani.docx')
doc.save(out)
print(f'Template saved: {out}')
